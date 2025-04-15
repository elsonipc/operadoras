import streamlit as st
import PyPDF2
import re
from docx import Document
from io import BytesIO
import datetime
import pandas as pd
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def extract_bo_data(pdf_file):
    """Extrai números de BO, IMEIs e Data/Hora do Fato de um arquivo PDF"""
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        
        for page in pdf_reader.pages:
            text += page.extract_text() or ""
        
        # Padrões para extração
        bo_pattern = r'N[º°oO]:\s*(\d+\s*/\s*\d+)'
        imei_pattern = r'IMEI\s*[2]?\s*[:=\s-]*\s*([\d\s]{15,20})'
        data_hora_pattern = r'Data/Hora\s*do\s*Fato\s*In[íi]cio\s*:\s*(\d{2}/\d{2}/\d{4}\s*\d{2}:\d{2})'
        
        # Busca as informações no texto
        bo_match = re.search(bo_pattern, text, re.IGNORECASE)
        
        # Encontra todos os IMEIs
        imeis = re.findall(imei_pattern, text, re.IGNORECASE)
        imeis_clean = [re.sub(r'[^\d]', '', imei)[:15] for imei in imeis if len(re.sub(r'[^\d]', '', imei)) >= 15]
        
        # Data/Hora do Fato
        data_hora_match = re.search(data_hora_pattern, text, re.IGNORECASE)
        data_hora = data_hora_match.group(1) if data_hora_match else "Não informado"
        
        return {
            "Nº BO": bo_match.group(1).replace(" ", "") if bo_match else "Não encontrado",
            "IMEI 1": imeis_clean[0] if len(imeis_clean) > 0 else "Não encontrado",
            "IMEI 2": imeis_clean[1] if len(imeis_clean) > 1 else "Não encontrado",
            "Data/Hora do Fato": data_hora,
            "Arquivo": pdf_file.name
        }
    except Exception as e:
        st.error(f"Erro ao extrair dados do PDF: {str(e)}")
        return None

def apply_document_styles(doc):
    """Aplica estilos consistentes ao documento"""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

def generate_word_document(bo_data, numero_oficio, data_inicio, data_fim):
    """Gera documento Word com os dados extraídos no formato do modelo"""
    try:
        doc = Document()
        apply_document_styles(doc)
        
        # Número do documento no topo direito
        p_num = doc.add_paragraph()
        p_num.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run = p_num.add_run("2500210467")
        run.font.size = Pt(9)
        
        # Cabeçalho institucional
        header = doc.add_paragraph()
        header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = header.add_run("GOVERNO DO ESTADO DO AMAZONAS\nPOLÍCIA CIVIL\n28º DISTRITO INTEGRADO DE POLÍCIA - MANAUS - AM")
        run.font.bold = True
        
        # Espaçamento
        doc.add_paragraph()
        
        # Número do ofício
        p_oficio = doc.add_paragraph()
        p_oficio.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p_oficio.add_run(f"Ofício nº {numero_oficio}/GDT/28ºDIP")
        run.font.bold = True
        run.font.size = Pt(12)
        
        # Data atual formatada
        meses = ["JANEIRO", "FEVEREIRO", "MARÇO", "ABRIL", "MAIO", "JUNHO",
                "JULHO", "AGOSTO", "SETEMBRO", "OUTUBRO", "NOVEMBRO", "DEZEMBRO"]
        hoje = datetime.datetime.now()
        data_formatada = f"Manaus/AM, {hoje.day} de {meses[hoje.month-1]} de {hoje.year}."
        p_data = doc.add_paragraph(data_formatada)
        
        # E-mail
        p_email = doc.add_paragraph()
        p_email.add_run("E-mail para resposta: 28dip@policiacivil.am.gov.br / elson_brito@policiacivil.am.gov.br")
        
        # Espaçamento
        doc.add_paragraph()
        
        # Saudação inicial
        doc.add_paragraph("Senhor Gerente:")
        
        # Corpo do texto
        texto_oficio = """
Cumprimentando V. Sa., tendo em vista a apuração de crimes ocorridos na área de circunscrição desta Delegacia de Polícia Distrital, conforme Boletins infra especificados, visando identificar e localizar a autoria delitiva, e em especial atenção ao disposto no Art. 17-B da Lei 9.613/1998, venho, através do presente, REQUISITAR, no prazo máximo de 10 dias, que nos informe os dados cadastrais das linhas infra especificadas.
"""
        doc.add_paragraph(texto_oficio)
        
        # Tabela com 3 colunas (BO, IMEI, Período)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Cabeçalho da tabela
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'BOLETIM DE OCORRÊNCIA'
        hdr_cells[1].text = 'IMEI'
        hdr_cells[2].text = 'PESQUISAR A PARTIR DE'
        
        # Adiciona dados à tabela
        periodo_pesquisa = f"{data_inicio} a {data_fim}"
        for item in bo_data:
            row_cells = table.add_row().cells
            row_cells[0].text = item['Nº BO']
            
            # IMEIs
            imei_text = item['IMEI 1']
            if item['IMEI 2'] != "Não encontrado":
                imei_text += f"\n{item['IMEI 2']}"
            row_cells[1].text = imei_text
            
            # Período de pesquisa
            row_cells[2].text = periodo_pesquisa
        
        # Texto final
        texto_final = """
Informamos que o descumprimento injustificado do presente expediente resultará na responsabilização por crime de desobediência, na forma do art. 330 do Código Penal Brasileiro, e para tanto aguardamos resposta no prazo de 10 dias a contar da data do recebimento deste.

Art. 17-B. A autoridade policial e o Ministério Público terão acesso, exclusivamente, aos dados cadastrais do investigado que informam qualificação pessoal, filiação e endereço, independentemente de autorização judicial, mantidos pela Justiça Eleitoral, pelas empresas telefônicas, pelas instituições financeiras, pelos provedores de internet e pelas administradoras de cartão de crédito.
"""
        doc.add_paragraph(texto_final)
        
        # Assinatura
        doc.add_paragraph("Atenciosamente,")
        doc.add_paragraph()
        doc.add_paragraph("_________________________________________")
        doc.add_paragraph("Dr. Maurício Ramos Viçoso Silva")
        doc.add_paragraph("Delegado(a) de Polícia")
        doc.add_paragraph()
        
        # Destinatário
        doc.add_paragraph("A")
        doc.add_paragraph("Ilmo. Sr. Gerente da Empresa VIVO / TIM / CLARO")
        
                
        return doc
    except Exception as e:
        st.error(f"Erro ao gerar documento Word: {str(e)}")
        return None

def main():
    # Custom header with title and signature
    st.markdown("""
    <style>
    .header {
        display: flex;
        justify-content: space-between;
        align-items: center;
        margin-bottom: 1rem;
    }
    .title {
        font-size: 24px;
        font-weight: bold;
    }
    .signature {
        font-size: 14px;
        font-style: italic;
        align-self: flex-end;
    }
    </style>
    """, unsafe_allow_html=True)

    st.markdown("""
    <div class="header">
        <div class="title">Gerador de Ofícios Operadoras - 28ºDIP</div>
        <div class="signature">Elaborado pelo IPC Elson Jr</div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("---")  # Horizontal line
    
    # Rest of the app
    with st.expander("🔧 Configurações do Ofício", expanded=True):
        col1, col2 = st.columns(2)
        with col1:
            numero_oficio = st.text_input("Número do Ofício (ex: 023/2025):", "023/2025")
        with col2:
            data_inicio_str = st.text_input("Data Início da Pesquisa (DD/MM/AAAA):", 
                                          (datetime.date.today() - datetime.timedelta(days=30)).strftime("%d/%m/%Y"))
            data_fim_str = st.text_input("Data Fim da Pesquisa (DD/MM/AAAA):", 
                                       datetime.date.today().strftime("%d/%m/%Y"))
            
            try:
                data_inicio = datetime.datetime.strptime(data_inicio_str, "%d/%m/%Y").strftime("%d/%m/%Y")
                data_fim = datetime.datetime.strptime(data_fim_str, "%d/%m/%Y").strftime("%d/%m/%Y")
            except ValueError:
                st.error("Formato de data inválido. Use DD/MM/AAAA")
                return
    
    uploaded_files = st.file_uploader("📤 Carregue os arquivos PDF dos BOs", type="pdf", accept_multiple_files=True)
    
    if uploaded_files:
        all_bo_data = []
        for file in uploaded_files:
            file_data = extract_bo_data(file)
            if file_data:
                all_bo_data.append(file_data)
        
        if all_bo_data:
            st.success(f"✅ {len(all_bo_data)} registro(s) extraído(s)")
            df = pd.DataFrame(all_bo_data)
            st.dataframe(df)
            
            doc = generate_word_document(all_bo_data, numero_oficio, data_inicio, data_fim)
            if doc:
                output = BytesIO()
                doc.save(output)
                st.download_button(
                    label="📥 Baixar Ofício em Word",
                    data=output.getvalue(),
                    file_name=f"Oficio_{numero_oficio.replace('/', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.warning("Nenhum dado foi extraído dos arquivos.")

if __name__ == "__main__":
    main()